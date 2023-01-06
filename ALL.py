# 作者：刘青松
# 地点：电子科技大学
# 时间：2021/10/10 9:36
# 联系方式：1965585326lqs@gmail.com
# 如果您对本代码有任何雅正意见，请不吝赐教，鄙人将感激不尽！

import sys
import re
import random
from PyQt5.Qt import *
from PyQt5 import QtCore, uic
from PyQt5.QtGui import QIcon
from PyQt5.QtWidgets import QApplication, QMessageBox
from docx import Document

timu1 = []  # 保存提取的选择题题目包括题号
timu2 = []  # 保存提取的填空题
timu3 = []  # 保存提取的判断题
score = 0  # 总分

############# 选择题 ############
document1 = Document(r"单项选择题/Tiku1_lqs.docx")

for line in document1.paragraphs:
    word = line.text
    data = re.findall(r'\d+\.(.*)\(', word)  # 每一行正则表达式规律:数字 . (全字符) (             ## .*表示多个字符，+表示重复一次或者多次，因为题号有一位数也有两位数，多的题库也有三位数
    if data:
        timu1.append(data[0])  # 此时timu[0]就是第一题，类推

L1 = random.sample(range(1, 10), 3)  # 选择题随机题号 1~10随机不重复3个数  timu[L1-1]是真正题号

t1 = timu1[L1[0]-1]
T1_tem = t1.replace("：", "(    )\n")  # 在中文：后换行   "A.", "\nA."
T1_b = T1_tem.replace("B.", "\nB.")
T1_c = T1_b.replace("C.", "\nC.")
T1 = T1_c.replace("D.", "\nD.")  # T1 表示第一题

t2 = timu1[L1[1]-1]
T2_tem = t2.replace("：", "(    )\n")
T2_b = T2_tem.replace("B.", "\nB.")
T2_c = T2_b.replace("C.", "\nC.")
T2 = T2_c.replace("D.", "\nD.")

t3 = timu1[L1[2]-1]
T3_tem = t3.replace("：", "(    )\n")
T3_b = T3_tem.replace("B.", "\nB.")
T3_c = T3_b.replace("C.", "\nC.")
T3 = T3_c.replace("D.", "\nD.")

############# 填空题 ############
document2 = Document(r"填空题/Tiku2_lqs.docx")

for line2 in document2.paragraphs:
    word2 = line2.text
    data2 = re.findall(r'\d+\、(.*)\。', word2)
    if data2:
        timu2.append(data2[0])

L2 = random.sample(range(1, 10), 3)

t4 = timu2[L2[0]-1]
T4 = t4 + '。'
t5 = timu2[L2[1]-1]
T5 = t5 + '。'
t6 = timu2[L2[2]-1]
T6 = t6 + '。'

############# 判断题 ############
document3 = Document(r"判断题/Tiku3_lqs.docx")

for line3 in document3.paragraphs:
    word3 = line3.text
    data3 = re.findall(r'\d+\、(.*)\。', word3)
    if data3:
        timu3.append(data3[0])

L3 = random.sample(range(1, 10), 3)

t7 = timu3[L3[0]-1]
T7 = t7 + '。'
t8 = timu3[L3[1]-1]
T8 = t8 + '。'
t9 = timu3[L3[2]-1]
T9 = t9 + '。'


class Login(QWidget):
    def __init__(self):
        super().__init__()
        self.ui = uic.loadUi('ui/LogIn4_anime.ui')  # ui/LogIn.ui
        self.ui.toolButton1.setIcon(QIcon('image/user3.jpg'))  # 添加user图标
        self.ui.setWindowFlags(QtCore.Qt.Window | QtCore.Qt.FramelessWindowHint)  # 无边框 再加上透明
        self.ui.lineEdit1.setFocus()  # 设置默认光标位置
        self.ui.setStyleSheet("#Form{border-image:url(image/orange.png);}")  # 设置背景颜色
        self.ui.button1.clicked.connect(self.handleCalc)  # 登录按钮点击后的行为
        self.ui.btn_close.clicked.connect(self.btn_close_clicked)
        self.ui.btn_min.clicked.connect(self.btn_min_clicked)
        self.ui.setAttribute(QtCore.Qt.WA_TranslucentBackground)  # 以下为登录界面优化，消除父窗背景

    def btn_close_clicked(self):  # 定义关闭窗口
        self.ui.close()

    def btn_min_clicked(self):  # 定义最小化窗口
        self.ui.showMinimized()

    def handleCalc(self):  # 点击会发生的事件
        if self.ui.radioButton1.isChecked() == 1:  # 教师端管理员账号与密码
            if self.ui.lineEdit1.text() == 'Admin':
                if self.ui.lineEdit2.text() == '000000':
                    QMessageBox.information(self.ui, '提示', '登录成功，点击继续')
                else:
                    QMessageBox.critical(self.ui, '提示', '用户名或密码错误')
            else:
                QMessageBox.critical(self.ui, '提示', '用户名或密码错误')

        elif self.ui.radioButton2.isChecked() == 1:  # 学生端账号与密码
            if self.ui.lineEdit1.text() == 'Admin':
                if self.ui.lineEdit2.text() == '111111':
                    choice = QMessageBox.information(self.ui, '提示', '登录成功，点击OK继续', QMessageBox.Ok | QMessageBox.Cancel)
                    if choice == QMessageBox.Ok:  # 判断是否点击了OK
                        self.ui.close()  # 关闭登录界面
                        dati.ui.show()  # 打开答题界面

                else:
                    QMessageBox.critical(self.ui, '提示', '用户名或密码错误')
            else:
                QMessageBox.critical(self.ui, '提示', '用户名或密码错误')
        else:
            QMessageBox.critical(self.ui, '提示', '请选择登录端')


class Answer(QWidget):
    def __init__(self):
        super().__init__()
        self.ui = uic.loadUi('ui/answer2.ui')
        self.ui.setWindowTitle('AnswerSheet by:Qingsong Liu UESTC')

        self.ui.textEdit.setFontFamily("微软雅黑")
        self.ui.textEdit.setFontPointSize(20)  # 字体大小
        # 选择题1~3
        self.ui.textEdit.setText('一、单项选择题(每题10分，共30分) \n')  # 覆盖原来内容
        self.ui.textEdit.append('1、%s\n' % T1)  # 增加内容
        self.ui.textEdit.append('2、%s\n' % T2)
        self.ui.textEdit.append('3、%s\n' % T3)
        # 填空题4~6
        self.ui.textEdit.append('二、填空题(每题10分，共30分) \n')
        self.ui.textEdit.append('4、%s\n' % T4)
        self.ui.textEdit.append('5、%s\n' % T5)
        self.ui.textEdit.append('6、%s\n' % T6)
        # 判断题7~9
        self.ui.textEdit.append('三、判断题(每题10分，共30分) \n')
        self.ui.textEdit.append('7、%s' % T7)
        self.ui.textEdit.setAlignment(Qt.AlignRight)
        self.ui.textEdit.append('(    )')
        self.ui.textEdit.setAlignment(Qt.AlignLeft)
        self.ui.textEdit.append('8、%s' % T8)
        self.ui.textEdit.setAlignment(Qt.AlignRight)
        self.ui.textEdit.append('(    )')
        self.ui.textEdit.setAlignment(Qt.AlignLeft)
        self.ui.textEdit.append('9、%s' % T9)
        self.ui.textEdit.setAlignment(Qt.AlignRight)
        self.ui.textEdit.append('(    )')
        self.ui.textEdit.setAlignment(Qt.AlignLeft)

        # 答题卡作答
        self.ui.pushButton_jiaojuan.clicked.connect(self.btn_jiaojuan_clicked)  # 交卷按钮

        self.ui.pushButton.clicked.connect(self.btn_t1_clicked)  # 答题卡第一题被点击
        self.ui.pushButton_2.clicked.connect(self.btn_t2_clicked)
        self.ui.pushButton_3.clicked.connect(self.btn_t3_clicked)
        self.ui.pushButton_4.clicked.connect(self.btn_t4_clicked)
        self.ui.pushButton_5.clicked.connect(self.btn_t5_clicked)
        self.ui.pushButton_6.clicked.connect(self.btn_t6_clicked)
        self.ui.pushButton_7.clicked.connect(self.btn_t7_clicked)
        self.ui.pushButton_8.clicked.connect(self.btn_t8_clicked)
        self.ui.pushButton_9.clicked.connect(self.btn_t9_clicked)

    def btn_t1_clicked(self):
        global score
        self.box = QMessageBox(QMessageBox.Question, '选择题1', '请选择您的答案')
        a = self.box.addButton('A', QMessageBox.YesRole)
        c = self.box.addButton('C', QMessageBox.NoRole)
        b = self.box.addButton('B', QMessageBox.AcceptRole)
        d = self.box.addButton('D', QMessageBox.RejectRole)
        self.box.setIcon(1)
        self.box.show()

        if self.box.clickedButton() == a:   # choice == QMessageBox.Ok:  # 判断是否点击了OK
            score = score + 10

    def btn_t2_clicked(self):
        global score
        self.box2 = QMessageBox(QMessageBox.Question, '选择题2', '请选择您的答案')
        a2 = self.box2.addButton('A', QMessageBox.YesRole)
        c2 = self.box2.addButton('C', QMessageBox.NoRole)
        b2 = self.box2.addButton('B', QMessageBox.AcceptRole)
        d2 = self.box2.addButton('D', QMessageBox.RejectRole)
        self.box2.setIcon(1)
        self.box2.show()

        if self.box2.clickedButton() == a2:
            score = score + 10

    def btn_t3_clicked(self):
        global score
        self.box3 = QMessageBox(QMessageBox.Question, '选择题2', '请选择您的答案')
        a3 = self.box3.addButton('A', QMessageBox.YesRole)
        c3 = self.box3.addButton('C', QMessageBox.NoRole)
        b3 = self.box3.addButton('B', QMessageBox.AcceptRole)
        d3 = self.box3.addButton('D', QMessageBox.RejectRole)
        self.box3.setIcon(1)
        self.box3.show()

        if self.box3.clickedButton() == a3:
            score = score + 10

    def btn_t4_clicked(self):
        QMessageBox.critical(self.ui, '提示', '填空题尚在开发')

    def btn_t5_clicked(self):
        QMessageBox.critical(self.ui, '提示', '填空题尚在开发')

    def btn_t6_clicked(self):
        QMessageBox.critical(self.ui, '提示', '填空题尚在开发')

    def btn_t7_clicked(self):
        global score
        self.box7 = QMessageBox(QMessageBox.Question, '判断题7', '请判断正误')
        a7 = self.box7.addButton('✔', QMessageBox.YesRole)
        b7 = self.box7.addButton('✘', QMessageBox.AcceptRole)
        self.box7.setIcon(1)
        self.box7.show()

    def btn_t8_clicked(self):
        global score
        self.box8 = QMessageBox(QMessageBox.Question, '判断题8', '请判断正误')
        a8 = self.box8.addButton('✔', QMessageBox.YesRole)
        b8 = self.box8.addButton('✘', QMessageBox.AcceptRole)
        self.box8.setIcon(1)
        self.box8.show()

    def btn_t9_clicked(self):
        global score
        self.box9 = QMessageBox(QMessageBox.Question, '判断题9', '请判断正误')
        a9 = self.box9.addButton('✔', QMessageBox.YesRole)
        b9 = self.box9.addButton('✘', QMessageBox.AcceptRole)
        self.box9.setIcon(1)
        self.box9.show()


    def btn_jiaojuan_clicked(self):
        global score
        QMessageBox.about(self.ui, '总分', '%d分' % score)


app = QApplication(sys.argv)
denglu = Login()
dati = Answer()
denglu.ui.show()
sys.exit(app.exec_())
